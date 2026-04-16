"""
File Type Detector — analyzes document content and detects what types
of data are present (tables, charts, text, images).

Returns a ContentProfile that drives the parallel extraction pipeline.

Improvements over v1:
  - Gemini Vision classifies images as chart/table/photo (not just filename keywords)
  - PDF table detection uses pymupdf find_tables() + improved text heuristics
  - PDF chart detection uses Gemini for large embedded images
"""
import base64
from dataclasses import dataclass, field
from typing import List, Optional
from pathlib import Path

from config.settings import settings
from utils.logger import get_logger

log = get_logger(__name__)


@dataclass
class ContentProfile:
    """Describes what content types were detected in a document."""
    file_path: str
    file_type: str  # pdf, docx, xlsx, csv, image

    has_tables: bool = False
    has_charts: bool = False
    has_text: bool = True
    has_images: bool = False

    page_count: int = 0
    table_page_indices: List[int] = field(default_factory=list)
    image_page_indices: List[int] = field(default_factory=list)

    raw_bytes: Optional[bytes] = field(default=None, repr=False)


def detect_content(loaded_file) -> ContentProfile:
    """Analyze a LoadedFile and return its ContentProfile."""
    ext = loaded_file.extension
    profile = ContentProfile(
        file_path=str(loaded_file.path),
        file_type=ext.lstrip("."),
        raw_bytes=loaded_file.raw_bytes,
    )

    if loaded_file.is_pdf:
        _detect_pdf_content(profile, loaded_file)
    elif loaded_file.is_docx:
        _detect_docx_content(profile, loaded_file)
    elif loaded_file.is_spreadsheet:
        _detect_spreadsheet_content(profile, loaded_file)
    elif loaded_file.is_image:
        _detect_image_content(profile, loaded_file)

    log.info(
        f"[bold]{loaded_file.filename}[/] → "
        f"tables={profile.has_tables}, charts={profile.has_charts}, "
        f"text={profile.has_text}, images={profile.has_images}, "
        f"pages={profile.page_count}"
    )
    return profile


# ──────────────────────────────────────────────────────────
# Gemini Vision classifier (shared by PDF + image detection)
# ──────────────────────────────────────────────────────────

def _classify_image_with_gemini(image_bytes: bytes, mime_type: str = "image/png") -> str:
    """Ask Gemini Vision to classify an image.

    Returns one of: 'chart', 'table', 'infographic', 'photo', 'unknown'
    """
    if not settings.gemini_api_key:
        return "unknown"

    try:
        from google import genai
        from utils.gemini_helper import call_gemini

        client = genai.Client(api_key=settings.gemini_api_key)
        b64_data = base64.b64encode(image_bytes).decode("utf-8")

        response = call_gemini(
            client=client,
            model=settings.gemini_model,
            contents=[{
                "role": "user",
                "parts": [
                    {"text": (
                        "Classify this image into EXACTLY ONE category. "
                        "Reply with ONLY the category name, nothing else.\n\n"
                        "Categories:\n"
                        "- chart (any graph, plot, bar chart, line chart, pie chart, histogram, scatter plot)\n"
                        "- table (structured data in rows and columns)\n"
                        "- infographic (educational poster, diagram with text)\n"
                        "- photo (photograph, screenshot, illustration)\n\n"
                        "Reply with ONE word only."
                    )},
                    {"inline_data": {"mime_type": mime_type, "data": b64_data}},
                ],
            }],
        )

        classification = response.text.strip().lower().split()[0]
        # Normalize
        if classification in ("chart", "graph", "plot", "histogram"):
            return "chart"
        elif classification in ("table", "spreadsheet"):
            return "table"
        elif classification in ("infographic", "poster", "diagram"):
            return "infographic"
        elif classification in ("photo", "photograph", "screenshot", "illustration"):
            return "photo"
        else:
            return classification

    except Exception as e:
        log.warning(f"Gemini image classification error: {e}")
        return "unknown"


# ──────────────────────────────────────────────────────────
# PDF detection
# ──────────────────────────────────────────────────────────

def _detect_pdf_content(profile: ContentProfile, loaded_file) -> None:
    """Detect content types inside a PDF.

    Improvements:
    - Uses pymupdf find_tables() for structural table detection
    - Uses improved text heuristics (lower threshold)
    - Uses Gemini Vision to classify large embedded images as charts
    """
    try:
        import fitz  # pymupdf

        doc = fitz.open(stream=loaded_file.raw_bytes, filetype="pdf")
        profile.page_count = len(doc)

        for page_idx, page in enumerate(doc):
            text = page.get_text("text")
            if text and len(text.strip()) > 20:
                profile.has_text = True

            # ── Table detection: pymupdf find_tables() ──
            try:
                found_tables = page.find_tables()
                if found_tables and len(found_tables.tables) > 0:
                    for ft in found_tables.tables:
                        data = ft.extract()
                        # Only count if it has at least 2 rows and 2 cols
                        if data and len(data) >= 2 and len(data[0]) >= 2:
                            profile.has_tables = True
                            if page_idx not in profile.table_page_indices:
                                profile.table_page_indices.append(page_idx)
            except Exception:
                pass  # find_tables not available in older pymupdf

            # ── Text-based table fallback ──
            if not profile.has_tables:
                lines = text.strip().split("\n") if text else []
                if len(lines) > 3:
                    table_like_lines = sum(
                        1 for line in lines
                        if (line.count("  ") >= 1 or "\t" in line or "|" in line)
                        and len(line.strip()) > 5
                    )
                    if table_like_lines >= 3:
                        profile.has_tables = True
                        if page_idx not in profile.table_page_indices:
                            profile.table_page_indices.append(page_idx)

            # ── Image & chart detection ──
            image_list = page.get_images(full=True)
            if image_list:
                profile.has_images = True
                profile.image_page_indices.append(page_idx)

                # Check each large image for chart content
                for img_info in image_list:
                    xref = img_info[0]
                    try:
                        base_image = doc.extract_image(xref)
                        w = base_image.get("width", 0)
                        h = base_image.get("height", 0)

                        if w >= 200 and h >= 200:
                            # Old heuristic: minimal text = likely chart
                            if len(text.strip()) < 200:
                                profile.has_charts = True
                            else:
                                # NEW: Use Gemini to classify large images
                                # even on pages with lots of text
                                ext = base_image.get("ext", "png")
                                img_class = _classify_image_with_gemini(
                                    base_image["image"],
                                    f"image/{ext}",
                                )
                                if img_class == "chart":
                                    profile.has_charts = True
                                    log.info(
                                        f"  Gemini classified PDF p{page_idx+1} "
                                        f"image as: [bold]{img_class}[/]"
                                    )
                    except Exception:
                        pass

        doc.close()
    except Exception as e:
        log.warning(f"PDF detection error: {e}")
        profile.has_text = True


# ──────────────────────────────────────────────────────────
# DOCX detection
# ──────────────────────────────────────────────────────────

def _detect_docx_content(profile: ContentProfile, loaded_file) -> None:
    """Detect content types inside a DOCX."""
    try:
        import io
        from docx import Document

        doc = Document(io.BytesIO(loaded_file.raw_bytes))

        if any(p.text.strip() for p in doc.paragraphs):
            profile.has_text = True

        if doc.tables:
            profile.has_tables = True

        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                profile.has_images = True
                break

        profile.page_count = 1
    except Exception as e:
        log.warning(f"DOCX detection error: {e}")
        profile.has_text = True


# ──────────────────────────────────────────────────────────
# Spreadsheet detection
# ──────────────────────────────────────────────────────────

def _detect_spreadsheet_content(profile: ContentProfile, loaded_file) -> None:
    """Detect content in spreadsheet files (XLSX, CSV)."""
    profile.has_tables = True
    profile.has_text = False
    profile.page_count = 1

    if loaded_file.extension in {".xlsx", ".xls"}:
        try:
            import io
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(loaded_file.raw_bytes), data_only=True)
            profile.page_count = len(wb.sheetnames)

            for sheet in wb.worksheets:
                if hasattr(sheet, '_charts') and sheet._charts:
                    profile.has_charts = True
                    break

            wb.close()
        except Exception as e:
            log.warning(f"XLSX detection error: {e}")


# ──────────────────────────────────────────────────────────
# Image detection (standalone images)
# ──────────────────────────────────────────────────────────

def _detect_image_content(profile: ContentProfile, loaded_file) -> None:
    """Image files — use Gemini Vision to detect charts/tables in images.

    Step 1: Check filename keywords (fast, no API call)
    Step 2: If no keyword match, ask Gemini to classify the image
    """
    profile.has_images = True
    profile.has_text = False
    profile.has_tables = False
    profile.page_count = 1

    name_lower = loaded_file.filename.lower()

    # Step 1: Filename keyword check (free, instant)
    chart_keywords = {"chart", "graph", "plot", "diagram", "histogram", "scatter"}
    if any(kw in name_lower for kw in chart_keywords):
        profile.has_charts = True
        log.info(f"  Chart detected by filename keyword: [bold]{loaded_file.filename}[/]")
        return

    # Step 2: Gemini Vision classification (handles all other images)
    ext = profile.file_type
    mime_map = {
        "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
        "bmp": "image/bmp", "tiff": "image/tiff", "gif": "image/gif",
        "webp": "image/webp",
    }
    mime_type = mime_map.get(ext, "image/png")

    img_class = _classify_image_with_gemini(loaded_file.raw_bytes, mime_type)
    log.info(f"  Gemini classified [bold]{loaded_file.filename}[/] as: [bold]{img_class}[/]")

    if img_class == "chart":
        profile.has_charts = True
    elif img_class == "table":
        # Image contains a table — flag as both chart (for extraction) and image
        profile.has_charts = True  # Route through chart→table pipeline
        log.info(f"  Table-in-image detected → routing through chart pipeline")
