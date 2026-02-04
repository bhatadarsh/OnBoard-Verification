
# from jd_intelligence.graph import jd_graph
# from resume_intelligence.graph import resume_graph
# from focus_area_selection.graph import stage3_graph
# from docx import Document


# def read_docx(path: str) -> str:
#     doc = Document(path)
#     return "\n".join(p.text for p in doc.paragraphs)


# # ---------------- JD STAGE ----------------
# jd_text = read_docx(
#     "/Users/sv-mac-313/Downloads/Job_Description_AI_Engineer.docx"
# )

# jd_state = {
#     "raw_jd": jd_text
# }

# jd_output = jd_graph.invoke(jd_state)


# # ---------------- RESUME STAGE ----------------
# resume_text = read_docx(
#     "/Users/sv-mac-313/Downloads/Matching_Resume_AI_Engineer.docx"
# )

# resume_state = {
#     "candidate_id": "CAND_001",
#     "raw_resume": resume_text,

#     # ---- JD-derived context ----
#     "role_context": jd_output["role_context"],
#     "skill_intelligence": jd_output["skill_intelligence"],
#     "competency_profile": jd_output["competency_profile"],
#     "interview_requirements": jd_output["interview_requirements"]
# }

# resume_output = resume_graph.invoke(resume_state)


# # ---------------- STAGE 2 RESULT ----------------
# print("\n=== STAGE 2 RESULT ===")
# print("FINAL SCORE:", resume_output["final_score"])
# print("SHORTLIST:", resume_output["shortlist_decision"])
# print("REASON:", resume_output["shortlist_reason"])


# # ---------------- STAGE 3 (ONLY IF SHORTLISTED) ----------------
# if resume_output["shortlist_decision"]:

#     # 🔑 Merge Stage 2 output into Stage 3 state
#     stage3_state = {
#         **resume_output,   # includes resume_claims, match_scores, etc.

#         # ---- JD context still needed ----
#         "interview_requirements": jd_output["interview_requirements"],
#         "skill_intelligence": jd_output["skill_intelligence"]
#     }

#     stage3_output = stage3_graph.invoke(stage3_state)
#     print("\n[DEBUG] Raw Stage 3 Output:")
#     print(stage3_output)

#     print("\n=== STAGE 3 – FINAL FOCUS AREAS ===")
#     for area in stage3_output.get("final_focus_areas", []):
#         print(area)

# else:
#     print("\nCandidate not shortlisted — Stage 3 skipped.")



from jd_intelligence.graph import jd_graph
from resume_intelligence.graph import resume_graph
from focus_area_selection.graph import stage3_graph
from interview_orchestration.graph import stage4_graph
from docx import Document


def read_docx(path: str) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


# ---------------- JD STAGE ----------------
jd_text = read_docx(
    "/Users/sv-mac-313/Downloads/Job_Description_AI_Engineer.docx"
)

jd_state = {"raw_jd": jd_text}
jd_output = jd_graph.invoke(jd_state)


# ---------------- RESUME STAGE ----------------
resume_text = read_docx(
    "/Users/sv-mac-313/Downloads/Matching_Resume_AI_Engineer.docx"
)

resume_state = {
    "candidate_id": "CAND_001",
    "raw_resume": resume_text,
    "role_context": jd_output["role_context"],
    "skill_intelligence": jd_output["skill_intelligence"],
    "competency_profile": jd_output["competency_profile"],
    "interview_requirements": jd_output["interview_requirements"],
}

resume_output = resume_graph.invoke(resume_state)


# ---------------- STAGE 2 RESULT ----------------
print("\n=== STAGE 2 RESULT ===")
print("FINAL SCORE:", resume_output["final_score"])
print("SHORTLIST:", resume_output["shortlist_decision"])
print("REASON:", resume_output["shortlist_reason"])


# ---------------- STAGE 3 ----------------
if not resume_output["shortlist_decision"]:
    print("\nCandidate not shortlisted — Stage 3 & 4 skipped.")
    exit()

stage3_state = {
    **resume_output,
    "interview_requirements": jd_output["interview_requirements"],
    "skill_intelligence": jd_output["skill_intelligence"],
}

stage3_output = stage3_graph.invoke(stage3_state)

print("\n=== STAGE 3 – FINAL FOCUS AREAS ===")
for area in stage3_output.get("final_focus_areas", []):
    print(area)


# ---------------- STAGE 4 (TEXT MODE INTERVIEW) ----------------
print("\n=== STAGE 4 – INTERVIEW (TEXT MODE) ===")

stage4_state = {
    "candidate_id": resume_state["candidate_id"],
    "final_focus_areas": stage3_output["final_focus_areas"],
}

SIMULATED_ANSWERS = [

    # 1️⃣ Very generic / LLM-like
    "A multi-agent workflow involves multiple autonomous agents working together to solve complex problems efficiently and effectively.",

    # 2️⃣ Copy-paste style repetition
    "A multi-agent workflow involves multiple autonomous agents working together to solve complex problems efficiently and effectively.",

    # 3️⃣ Over-verbose + irrelevant
    "In modern AI systems, multi-agent architectures are extremely important across many domains including finance, healthcare, logistics, education, and many others, where scalability, modularity, and robustness are critical considerations for enterprise-grade deployments.",

    # 4️⃣ Sudden topic drift
    "JWT authentication uses access tokens and refresh tokens and is commonly used in REST APIs."
]


answer_index = 0

# 🔑 Initialize interview
state = stage4_graph.invoke(stage4_state, config={"recursion_limit": 200})

# 🔁 Controlled event-driven loop
while state["interview_status"] != "COMPLETED":

    if state["interview_status"] == "WAITING_FOR_ANSWER":
        print(f"\nQ: {state['current_question']}")

        state["simulated_answer"] = SIMULATED_ANSWERS[
            min(answer_index, len(SIMULATED_ANSWERS) - 1)
        ]
        answer_index += 1

        state = stage4_graph.invoke(state, config={"recursion_limit": 200})

    else:
        # Allow graph to reach WAITING_FOR_ANSWER
        state = stage4_graph.invoke(state, config={"recursion_limit": 200})

print("\n=== INTERVIEW TRACE ===")
for turn in state["interview_trace"]:
    print(
        f"[{turn['topic']} | {turn['question_type']} #{turn['followup_index']}] "
        f"Q: {turn['question']} | A: {turn['answer_text']}"
    )

print("\n=== CHEATING ANALYSIS ===")
print("Cheating score:", state.get("cheating_score"))
print("Cheating events:")
for event in state.get("cheating_events", []):
    print(event)
