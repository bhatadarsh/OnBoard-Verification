from docx import Document

# 1. Create Dummy JD
jd = Document()
jd.add_heading('Senior Software Engineer', 0)
jd.add_paragraph('We are looking for a Python expert with Azure experience.')
jd.add_heading('Responsibilities', level=1)
jd.add_paragraph('- Design and implement REST APIs using Python.')
jd.add_paragraph('- Manage Azure Blob Storage integrations.')
jd.add_heading('Requirements', level=1)
jd.add_paragraph('- 5+ years of Python experience.')
jd.add_paragraph('- Experience with Azure Speech Services.')
jd.save('dummy_jd.docx')
print("✅ Created dummy_jd.docx")

# 2. Create Dummy Resume
res = Document()
res.add_heading('Test Candidate', 0)
res.add_paragraph('Python Developer | Azure Specialist')
res.add_heading('Experience', level=1)
res.add_paragraph('Senior Python Developer at Tech Corp.')
res.add_paragraph('- Built scalable APIs using FastAPI.')
res.add_paragraph('- Integrated Azure Cognitive Services.')
res.add_heading('Skills', level=1)
res.add_paragraph('Python, Azure, Streamlit, SQL.')
res.save('dummy_resume.docx')
print("✅ Created dummy_resume.docx")
