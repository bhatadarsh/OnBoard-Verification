"""
Ingestion Tools - Document text extraction utilities used by the ingestion node.
Handles massive concurrent OCR (via Groq Vision), PDF, DOCX, and plain text extraction.
"""
import os
import base64
import tempfile
import asyncio
from typing import Optional
from contextlib import contextmanager
from cryptography.fernet import Fernet

try:
    from pypdf import PdfReader
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from app.core.config import config
from app.services.llm_service import llm_service


@contextmanager
def decrypted_tempfile(file_path: str):
    """Decrypt a file into a temporary location and cleanly shred it after use."""
    if not os.path.exists(file_path):
        yield ""
        return
        
    ext = os.path.splitext(file_path)[1].lower()
    tmp_path = ""
    try:
        fernet = Fernet(config.FERNET_KEY.encode('utf-8'))
        with open(file_path, 'rb') as f:
            encrypted = f.read()
            
        try:
            data = fernet.decrypt(encrypted)
        except Exception:
            data = encrypted
            
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
            
        yield tmp_path
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


async def extract_text_from_file(file_path: str) -> str:
    """Extract text from any supported document format asynchronously.
    
    Supports: images (OCR via Groq Vision), PDF, DOCX, plain text.
    """
    if not os.path.exists(file_path):
        return ""
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif', '.webp']:
        return await llm_service.extract_from_image_vision_async(file_path)
    elif ext == '.pdf':
        return await _extract_from_pdf(file_path)
    elif ext == '.docx':
        return _extract_from_docx(file_path)
    elif ext in ['.txt', '.text', '.md']:
        return _extract_from_text(file_path)
    elif ext in ['.mp3', '.wav', '.m4a', '.ogg', '.flac']:
        return f"[AUDIO_FILE:{file_path}]"
    else:
        return _extract_from_text(file_path)


def is_audio_file(file_path: str) -> bool:
    """Check if file is audio."""
    ext = os.path.splitext(file_path)[1].lower()
    return ext in ['.mp3', '.wav', '.m4a', '.ogg', '.flac']


def _extract_from_text(file_path: str) -> str:
    """Read plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        return f"Error reading text: {e}"


async def _extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF. Falls back to Multi-Page Vision OCR for scanned PDFs."""
    if PdfReader is None:
        return "[PDF extraction not available - install PyPDF2]"
    
    try:
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        result = "\n".join(text_parts)
        
        if len(result.strip()) < 50:
            return await _extract_from_scanned_pdf(file_path)
        
        return result
    except Exception as e:
        return f"Error reading PDF: {e}"


async def _extract_from_scanned_pdf(file_path: str) -> str:
    """Extract text from scanned PDF using Concurrent Vision API mapped across multiple pages."""
    try:
        from pdf2image import convert_from_path
        
        # Pull ALL pages instead of strictly first
        images = convert_from_path(file_path)
        if not images:
            return "[Could not convert PDF to image]"
        
        tmp_files = []
        for img in images:
            tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            img.save(tmp.name, 'PNG')
            tmp_files.append(tmp.name)
            
        print(f"  Scanned PDF detected. Triggering internal parallel Vision OCR on {len(tmp_files)} pages.")
        
        tasks = [llm_service.extract_from_image_vision_async(p) for p in tmp_files]
        raw_results = await asyncio.gather(*tasks, return_exceptions=True)
        
        results = []
        for i, res in enumerate(raw_results):
            if isinstance(res, Exception):
                print(f"  [ERROR] Vision OCR failed on page {i}: {res}")
                results.append(f"[Page {i} failed to OCR]")
            else:
                results.append(str(res))
        
        for tmp_path in tmp_files:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
                
        return "\n--- Page Break ---\n".join(results)
    except Exception as e:
        return f"Error OCR on scanned PDF: {e}"


def _extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX file including tables."""
    if DocxDocument is None:
        return "[DOCX extraction not available - install python-docx]"
    
    try:
        doc = DocxDocument(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        return "\n".join(text_parts)
    except Exception as e:
        return f"Error reading DOCX: {e}"

async def verify_document_signature(file_path: str) -> str:
    """Verify signature for an image or PDF file using Vision AI."""
    import json
    if not os.path.exists(file_path):
        return '{"signature_detected": "No", "signature_details": "File not found"}'
        
    ext = os.path.splitext(file_path)[1].lower()
    if ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif', '.webp']:
        return await llm_service.verify_signature_vision_async(file_path)
    elif ext == '.pdf':
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(file_path)
            if not images:
                return '{"signature_detected": "No", "signature_details": "Failed to extract pdf pages"}'
            
            # Check first and last page (usually where signatures are)
            pages_to_check = [images[0]]
            if len(images) > 1:
                pages_to_check.append(images[-1])
                    
            last_res = "{}"
            for i, img in enumerate(pages_to_check):
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                    img.save(tmp.name, 'PNG')
                    tmp_path = tmp.name
                    
                res = await llm_service.verify_signature_vision_async(tmp_path)
                os.unlink(tmp_path)
                last_res = res
                
                try:
                    res_dict = json.loads(res)
                    if res_dict.get("signature_detected", "").lower() in ["yes", "true"]:
                        return res
                except:
                    pass
            # if none found, return the last result
            return last_res
        except Exception as e:
            return f'{{"signature_detected": "No", "signature_details": "PDF vision error: {e}"}}'
    return '{"signature_detected": "No", "signature_details": "Unsupported format for signature check"}'
